from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches


def export_consolidated_word(rows: List[Dict[str, Any]], title: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "CODIGO"
    hdr_cells[1].text = "PRODUTO"
    hdr_cells[2].text = "QUANTIDADE"
    hdr_cells[3].text = "UN"

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = str(row.get("code"))
        cells[1].text = str(row.get("name"))
        cells[2].text = str(row.get("quantity"))
        cells[3].text = str(row.get("unit"))

    from io import BytesIO
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


